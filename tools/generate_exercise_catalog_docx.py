import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CATALOG = ROOT / "artifacts" / "exercise-catalog" / "catalog.json"
OUTPUT = ROOT / "artifacts" / "当前可用训练动作目录.docx"


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = props.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        props.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size, bold=False, color="202326"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def build():
    rows = json.loads(CATALOG.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.75)
    section.left_margin = Cm(0.85)
    section.right_margin = Cm(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("形域 · 当前可用训练动作目录"), 8, color="6E757C")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("第 "), 8, color="7A8086")
    add_page_field(footer)
    set_font(footer.add_run(" 页"), 8, color="7A8086")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(62)
    title.paragraph_format.space_after = Pt(16)
    set_font(title.add_run("当前可用训练动作目录"), 28, bold=True, color="17191B")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(subtitle.add_run(f"共 {len(rows)} 个动作 · 图片、中文名称与稳定编号"), 13, color="E66A22")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Cm(5)
    note.paragraph_format.right_indent = Cm(5)
    note.paragraph_format.line_spacing = 1.5
    set_font(
        note.add_run("编号仅用于目录核对与历史数据定位；应用界面不显示编号。目录由当前可选择动作自动生成。"),
        11,
        color="555B61",
    )
    document.add_page_break()

    per_page = 10
    for page_start in range(0, len(rows), per_page):
        page_rows = rows[page_start : page_start + per_page]
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(6)
        set_font(
            heading.add_run(
                f"动作 {page_start + 1}–{page_start + len(page_rows)} / {len(rows)}"
            ),
            10,
            bold=True,
            color="555B61",
        )
        table = document.add_table(rows=2, cols=5)
        table.autofit = False
        table.allow_autofit = False
        for table_row in table.rows:
            table_row.height = Cm(7.5)
            table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for index, item in enumerate(page_rows):
            cell = table.cell(index // 5, index % 5)
            cell.width = Cm(5.45)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "F4F5F6")
            set_cell_margins(cell)
            image = cell.paragraphs[0]
            image.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image.paragraph_format.space_after = Pt(2)
            asset = ROOT / "mobile" / item["asset"]
            image.add_run().add_picture(str(asset), height=Cm(4.7))
            number = cell.add_paragraph()
            number.alignment = WD_ALIGN_PARAGRAPH.CENTER
            number.paragraph_format.space_after = Pt(0)
            set_font(number.add_run(f"编号 {item['number']}"), 8, bold=True, color="E66A22")
            name = cell.add_paragraph()
            name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name.paragraph_format.space_before = Pt(0)
            set_font(name.add_run(item["name"]), 10, bold=True, color="202326")
        for empty_index in range(len(page_rows), per_page):
            cell = table.cell(empty_index // 5, empty_index % 5)
            set_cell_shading(cell, "FFFFFF")
        if page_start + per_page < len(rows):
            document.add_page_break()

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
